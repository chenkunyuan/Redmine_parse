"""Generate Redmine Parse User Guide docx.

Matches the Hydrant Log Spec document format:
  - Cover page (no header/footer)
  - Sercomm confidential footer on subsequent pages
  - Proper TOC field
  - A4 page, matching margins
"""
import sys
sys.stdout.reconfigure(encoding="utf-8")

from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_TAB_ALIGNMENT, WD_TAB_LEADER
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml, OxmlElement

doc = Document()

# --- Page setup: match reference margins ---
for section in doc.sections:
    section.page_width = Cm(21.0)
    section.page_height = Cm(29.7)
    section.left_margin = Cm(1.42)
    section.right_margin = Cm(0.95)
    section.top_margin = Cm(1.9)
    section.bottom_margin = Cm(1.9)
    section.different_first_page_header_footer = True

# --- Styles ---
style = doc.styles["Normal"]
style.font.name = "Arial"
style.font.size = Pt(10)
style.paragraph_format.space_after = Pt(6)
style.paragraph_format.line_spacing = 1.15

for lvl, (size, color) in enumerate([(16, "1F3864"), (13, "2E75B6"), (11, "2E75B6")], 1):
    h = doc.styles[f"Heading {lvl}"]
    h.font.name = "Arial"
    h.font.size = Pt(size)
    h.font.bold = True
    h.font.color.rgb = RGBColor.from_string(color)
    h.paragraph_format.space_before = Pt(12)
    h.paragraph_format.space_after = Pt(6)


def code(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(8)
    run = p.add_run(text)
    run.font.name = "Consolas"
    run.font.size = Pt(9)
    return p


def add_table(headers, rows, col_widths=None):
    table = doc.add_table(rows=len(rows) + 1, cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for j, h in enumerate(headers):
        cell = table.cell(0, j)
        cell.text = h
        for p in cell.paragraphs:
            for r in p.runs:
                r.font.size = Pt(9)
                r.font.name = "Arial"
                r.font.bold = True
    for i, row in enumerate(rows, 1):
        for j, val in enumerate(row):
            cell = table.cell(i, j)
            cell.text = val
            for p in cell.paragraphs:
                for r in p.runs:
                    r.font.size = Pt(9)
                    r.font.name = "Arial"
    if col_widths:
        for i, w in enumerate(col_widths):
            for row in table.rows:
                row.cells[i].width = Cm(w)
    return table


def bullet(text):
    p = doc.add_paragraph(text, style="List Bullet")
    for r in p.runs:
        r.font.name = "Arial"


def add_toc_field():
    """Insert a proper Word TOC field via XML."""
    p = doc.add_paragraph()
    run = p.add_run()
    fldChar_begin = OxmlElement("w:fldChar")
    fldChar_begin.set(qn("w:fldCharType"), "begin")
    run._r.append(fldChar_begin)

    run2 = p.add_run()
    instrText = OxmlElement("w:instrText")
    instrText.set(qn("xml:space"), "preserve")
    instrText.text = ' TOC \\o "1-3" \\h \\z '
    run2._r.append(instrText)

    run3 = p.add_run()
    fldChar_separate = OxmlElement("w:fldChar")
    fldChar_separate.set(qn("w:fldCharType"), "separate")
    run3._r.append(fldChar_separate)

    run4 = p.add_run("[ Right-click and select \"Update Field\" to generate ]")
    run4.font.size = Pt(9)
    run4.font.name = "Arial"
    run4.font.color.rgb = RGBColor(128, 128, 128)

    run5 = p.add_run()
    fldChar_end = OxmlElement("w:fldChar")
    fldChar_end.set(qn("w:fldCharType"), "end")
    run5._r.append(fldChar_end)


def add_footer(section):
    """Add Sercomm confidential footer."""
    footer = section.footer
    footer.is_linked_to_previous = False
    p = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT

    parts = [
        ("Sercomm Proprietary. Copyright ", "Calibri"),
        ("©", "Calibri"),
        (" 2026. All rights reserved. ", "Calibri"),
        ("All specifications are subject to change without notice.", "Calibri"),
    ]
    for text, font_name in parts:
        run = p.add_run(text)
        run.font.name = font_name
        run.font.size = Pt(7)


def add_header_line(section):
    """Add a thin line in header (matching reference style)."""
    header = section.header
    header.is_linked_to_previous = False
    p = header.paragraphs[0] if header.paragraphs else header.add_paragraph()
    # Bottom border as separator line
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "4")
    bottom.set(qn("w:space"), "4")
    bottom.set(qn("w:color"), "808080")
    pBdr.append(bottom)
    pPr.append(pBdr)


# ============ COVER PAGE ============
for _ in range(6):
    doc.add_paragraph("")

title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title.add_run("Redmine Parse")
run.font.size = Pt(28)
run.font.bold = True
run.font.name = "Arial"
run.font.color.rgb = RGBColor.from_string("1F3864")

subtitle = doc.add_paragraph()
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = subtitle.add_run("User Guide")
run.font.size = Pt(20)
run.font.name = "Arial"
run.font.color.rgb = RGBColor.from_string("2E75B6")

doc.add_paragraph("")
doc.add_paragraph("")

ver = doc.add_paragraph()
ver.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = ver.add_run("V1.0")
run.font.size = Pt(14)
run.font.bold = True
run.font.name = "Arial"

for _ in range(4):
    doc.add_paragraph("")

for text in [
    "Copyright 2026 by SerComm Corp.",
    "All rights reserved.",
    "This document contains information of a proprietary nature.",
    "ALL INFORMATION CONTAINED HEREIN SHALL BE KEPT IN STRICT CONFIDENTIAL.",
]:
    p = doc.add_paragraph(text)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for r in p.runs:
        r.font.size = Pt(8)
        r.font.name = "Arial"

doc.add_paragraph("")
doc.add_paragraph("")

for text in [
    "SerComm (Suzhou) R&D Center",
    "5F, No.26, Xinghai Street, Suzhou Industrial Park",
    "Jiangsu, China",
    "Tel: 86-512-67612332",
    "Fax: 86-512-67622332",
    "http://www.sercomm.com",
]:
    p = doc.add_paragraph(text)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for r in p.runs:
        r.font.size = Pt(8)
        r.font.name = "Arial"

# --- Add section break after cover to enable different first page ---
# The first section already has different_first_page_header_footer = True,
# so cover page gets no header/footer.
# We add a new section for page 2+ which will have headers/footers.
doc.add_section()
sec2 = doc.sections[1]
sec2.page_width = Cm(21.0)
sec2.page_height = Cm(29.7)
sec2.left_margin = Cm(1.42)
sec2.right_margin = Cm(0.95)
sec2.top_margin = Cm(1.9)
sec2.bottom_margin = Cm(1.9)
sec2.different_first_page_header_footer = False

add_header_line(sec2)
add_footer(sec2)

# ============ REVISION HISTORY ============
doc.add_heading("Revision History", level=1)
add_table(
    ["Version", "Date", "Author", "Section", "Description"],
    [["V1.0", "2026-05-25", "Cooper Chen", "All", "Initial version"]],
    col_widths=[2, 2.5, 2.5, 1.5, 8],
)
doc.add_paragraph("")
doc.add_page_break()

# ============ TABLE OF CONTENTS ============
doc.add_heading("Table of Contents", level=1)
add_toc_field()
doc.add_page_break()

# ============ 1. INTRODUCTION ============
doc.add_heading("1. Introduction", level=1)

doc.add_heading("1.1 Purpose", level=2)
doc.add_paragraph(
    "Redmine Parse is a command-line interface (CLI) tool for interacting with the Redmine "
    "issue tracking system via its REST API. It allows users to list projects, browse issues "
    "within projects, view issue details, create and update issues, and export issue data to "
    "Excel spreadsheets for reporting and analysis."
)
doc.add_paragraph(
    "The tool is designed primarily for the HYDRANT Field Issue Tracking workflow at SerComm, "
    "where field-returned devices are tracked as Redmine issues with associated metadata such "
    "as KDDI control numbers, serial numbers, and failure analysis reports."
)

doc.add_heading("1.2 System Requirements", level=2)
doc.add_paragraph("The following software is required:")
for item in [
    "Python 3.9 or later",
    "pip package manager",
    "Network access to the Redmine server (https://redmine.sercomm.co.jp)",
    "A valid Redmine API key with appropriate permissions",
]:
    bullet(item)

doc.add_page_break()

# ============ 2. INSTALLATION ============
doc.add_heading("2. Installation", level=1)

doc.add_heading("2.1 Prerequisites", level=2)
doc.add_paragraph("Ensure Python 3.9+ and pip are installed:")
code("python --version\npip --version")

doc.add_heading("2.2 Installation Steps", level=2)
doc.add_paragraph("1. Navigate to the project directory:")
code("cd Redmine_parse")
doc.add_paragraph("2. Install the package in development mode:")
code("pip install -e .")
doc.add_paragraph(
    "This installs the redmine CLI command along with its dependencies: "
    "click, python-redmine, requests, and openpyxl."
)
doc.add_paragraph("3. Verify the installation:")
code("redmine --help")
doc.add_paragraph(
    "You should see the top-level command group with subcommands: projects and issues."
)

doc.add_page_break()

# ============ 3. CONFIGURATION ============
doc.add_heading("3. Configuration", level=1)

doc.add_heading("3.1 Environment Variables", level=2)
doc.add_paragraph(
    "The tool uses two environment variables for authentication. Set them in your shell "
    "profile or before each session:"
)

add_table(
    ["Variable", "Required", "Description"],
    [
        ["REDMINE_URL", "No", "Redmine server URL (default: https://redmine.sercomm.co.jp)"],
        ["REDMINE_API_KEY", "Yes", "Your Redmine API access key"],
    ],
    col_widths=[3.5, 1.5, 10.5],
)

doc.add_paragraph("")
doc.add_paragraph("Example (Linux/macOS):")
code('export REDMINE_API_KEY="your-api-key"')
doc.add_paragraph("Example (Windows PowerShell):")
code('$env:REDMINE_API_KEY="your-api-key"')

doc.add_heading("3.2 Command-line Options", level=2)
doc.add_paragraph(
    "Both URL and API key can also be passed directly as command-line flags, which take "
    "precedence over environment variables:"
)

add_table(
    ["Flag", "Description", "Default"],
    [
        ["--url", "Redmine server URL", "REDMINE_URL env or https://redmine.sercomm.co.jp"],
        ["--key", "Redmine API key", "REDMINE_API_KEY env"],
    ],
    col_widths=[2, 4.5, 9],
)

doc.add_page_break()

# ============ 4. USAGE ============
doc.add_heading("4. Usage", level=1)
doc.add_paragraph(
    "All commands output JSON to stdout, making them suitable for piping to other tools "
    "(e.g., jq, python -m json.tool)."
)

doc.add_heading("4.1 List Projects", level=2)
doc.add_paragraph("Retrieve a list of all visible projects:")
code("redmine projects list")
doc.add_paragraph(
    "Returns a JSON array of project objects, each containing id, name, identifier, "
    "description, status, created_on, updated_on, and custom_fields."
)

doc.add_heading("4.2 List Issues", level=2)
doc.add_paragraph("List all issues within a specific project:")
code("redmine issues list --project-id 141")
doc.add_paragraph(
    "The --project-id (or -p) argument is required. Replace 141 with the target project ID. "
    "Each issue object includes: id, subject, tracker, status, priority, author, assigned_to, "
    "fixed_version, custom_fields, created_on, and updated_on."
)

doc.add_heading("4.3 View Issue Details", level=2)
doc.add_paragraph("Retrieve full details of a single issue:")
code("redmine issues show 14252")
doc.add_paragraph(
    "Returns a JSON object with all issue fields including the full description, journals, "
    "attachments, and relations."
)

doc.add_heading("4.4 Create and Update Issues", level=2)
doc.add_paragraph("Create a new issue:")
code('redmine issues create --project-id 141 --subject "New issue" --description "Details"')
doc.add_paragraph("Update an existing issue:")
code('redmine issues update 14252 --notes "Added analysis note" --status-id 2')
doc.add_paragraph("Available options for create and update:")

add_table(
    ["Option", "Description", "Notes"],
    [
        ["--project-id, -p", "Project ID", "Required for create"],
        ["--subject, -s", "Issue subject/title", "Required for create"],
        ["--description, -d", "Issue description text", "Optional"],
        ["--tracker-id", "Tracker ID", "Numeric"],
        ["--status-id", "Status ID", "Numeric"],
        ["--priority-id", "Priority ID", "Numeric"],
        ["--assigned-to-id", "Assignee user ID", "Numeric"],
        ["--notes", "Comment/note text", "Update only"],
    ],
    col_widths=[4, 5, 6.5],
)

doc.add_page_break()

doc.add_heading("4.5 Export to Excel", level=2)
doc.add_paragraph(
    "The project includes a dedicated script for exporting HYDRANT Field Issue Tracking "
    "data to Excel:"
)
code("python export_hydrant_issues.py")
doc.add_paragraph(
    "The script produces an Excel file (HYDRANT_Field_Issues.xlsx) with the following columns:"
)

add_table(
    ["Column", "Description"],
    [
        ["Issue", "Tracker name + issue ID (e.g. \"Returned Item_K #14252\"), hyperlinked to Redmine"],
        ["Subject", "Issue subject/title"],
        ["Priority", "Issue priority (e.g., Normal, High)"],
        ["Status", "Current status (e.g., New, In Progress, Resolved)"],
        ["Category", "Auto-classified category based on subject/description analysis"],
        ["Reporter", "Fixed value: Field"],
        ["Report Time", "Issue creation month in YYYY-MM format"],
    ],
    col_widths=[3, 12.5],
)

doc.add_page_break()

# ============ 5. EXCEL EXPORT TOOL ============
doc.add_heading("5. Excel Export Tool", level=1)

doc.add_heading("5.1 Export Script", level=2)
doc.add_paragraph("The export script is located at:")
code("Redmine_parse/export_hydrant_issues.py")
doc.add_paragraph("Usage:")
code(
    "# Incremental mode (default) -- only add new issues\n"
    "python export_hydrant_issues.py\n\n"
    "# Full rebuild -- regenerate all rows\n"
    "python export_hydrant_issues.py --full"
)

doc.add_heading("5.2 Incremental Update", level=2)
doc.add_paragraph(
    "By default, the script runs in incremental mode. It reads the existing Excel file, "
    "extracts all issue IDs already present in column A, then fetches the current issue list "
    "from Redmine. Only new issues (those with IDs not already in the sheet) are appended as "
    "new rows."
)
doc.add_paragraph("Benefits of incremental mode:")
for b in [
    "Preserves manual edits or annotations made to existing rows",
    "Faster -- only fetches and appends new data",
    "Safe for daily updates -- existing data is never overwritten",
    "New issues are appended at the bottom of the sheet",
]:
    bullet(b)

doc.add_paragraph(
    "Use --full mode when you need to rebuild the entire spreadsheet from scratch "
    "(e.g., after changing the column layout, category rules, or fixing data inconsistencies)."
)

doc.add_heading("5.3 Category Classification", level=2)
doc.add_paragraph(
    "The script automatically classifies each issue into one of the following categories "
    "based on keyword analysis of the subject line and description:"
)

add_table(
    ["Category", "Description"],
    [
        ["WAN/Connection", "Internet/network connection issues, ONU/optical, PPPoE, DHCP"],
        ["WIFI", "Wireless LAN, SSID, signal strength, association/disassociation"],
        ["Power/Reboot", "Power failure, unexpected reboots, shutdown"],
        ["LED/Lamp", "LED indicator anomalies, blinking patterns"],
        ["VoIP/TEL", "Telephone, SIP registration, voice quality"],
        ["LAN", "Ethernet port, cable issues"],
        ["FW/Filter", "Firewall, packet filtering, blocking"],
        ["GUI/Setting", "Web UI, configuration settings"],
        ["HW Defect", "Physical damage, hardware malfunction, burn marks"],
        ["Other", "Issues that do not match any specific category"],
    ],
    col_widths=[4, 11.5],
)

doc.add_page_break()

# ============ 6. REFERENCE ============
doc.add_heading("6. Reference", level=1)

doc.add_heading("6.1 Commands Reference", level=2)

add_table(
    ["Command", "Description", "Example"],
    [
        ["projects list", "List all visible projects", "redmine projects list"],
        ["issues list", "List issues in a project", "redmine issues list -p 141"],
        ["issues show", "View single issue details", "redmine issues show 14252"],
        ["issues create", "Create a new issue", 'redmine issues create -p 141 -s "Bug"'],
        ["issues update", "Update an existing issue", "redmine issues update 14252 --status-id 3"],
        ["export_hydrant_issues.py", "Export to Excel (incremental)", "python export_hydrant_issues.py"],
        ["export_hydrant_issues.py --full", "Export to Excel (full rebuild)", "python export_hydrant_issues.py --full"],
    ],
    col_widths=[4.5, 5, 6],
)

doc.add_paragraph("")

doc.add_heading("6.2 Redmine REST API", level=2)
doc.add_paragraph(
    "This tool uses the python-redmine library, which wraps the Redmine REST API. "
    "The API documentation is available at:"
)
code("https://www.redmine.org/projects/redmine/wiki/Rest_api")

doc.add_paragraph("The following API endpoints are used internally:")
for ep in [
    "GET /projects.json -- list all projects",
    "GET /issues.json?project_id=X -- list issues in a project",
    "GET /issues/X.json -- get issue details",
    "POST /issues.json -- create an issue",
    "PUT /issues/X.json -- update an issue",
]:
    bullet(ep)

# Save
output_path = "D:/Claude/Projects/Redmine-Parse/Redmine_parse/Redmine_Parse_User_Guide.docx"
doc.save(output_path)
print(f"Saved: {output_path}")
